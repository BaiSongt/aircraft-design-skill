#!/usr/bin/env python3
import sys
try:
    from docx import Document
    
    doc_path = '/Users/baisongtao/mycode/aircraft-design-skill/docs/飞机设计文件.docx'
    doc = Document(doc_path)
    
    # 提取所有段落文本
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # 提取所有表格内容
    for table in doc.tables:
        full_text.append("\n[表格开始]")
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            full_text.append(row_text)
        full_text.append("[表格结束]\n")
    
    # 输出到文件
    output_path = '/Users/baisongtao/mycode/aircraft-design-skill/docs/飞机设计文件.txt'
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(full_text))
    
    print(f"文档内容已提取到: {output_path}")
    print(f"共提取 {len(full_text)} 行内容")
    
except ImportError:
    print("错误: 未安装 python-docx 模块")
    sys.exit(1)
except Exception as e:
    print(f"错误: {e}")
    sys.exit(1)
